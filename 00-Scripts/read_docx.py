#!/usr/bin/env python3
"""
读取Word文档内容
"""

import sys
import os
from pathlib import Path

# 尝试导入python-docx
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def read_docx_file(filepath):
    """读取Word文档内容"""
    if not HAS_DOCX:
        print("错误: 需要安装python-docx库")
        print("请运行: pip install python-docx")
        return None
    
    try:
        doc = Document(filepath)
        content = []
        
        # 读取所有段落
        for para in doc.paragraphs:
            if para.text.strip():  # 跳过空行
                content.append(para.text)
        
        # 读取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
    
    except Exception as e:
        print(f"读取文档时出错: {e}")
        return None

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python read_docx.py <docx文件路径>")
        return
    
    filepath = sys.argv[1]
    
    if not os.path.exists(filepath):
        print(f"文件不存在: {filepath}")
        return
    
    print(f"正在读取: {filepath}")
    print("=" * 60)
    
    content = read_docx_file(filepath)
    
    if content:
        print(content)
        print("=" * 60)
        print(f"文档读取完成，共 {len(content.splitlines())} 行")
    else:
        print("无法读取文档内容")

if __name__ == "__main__":
    main()